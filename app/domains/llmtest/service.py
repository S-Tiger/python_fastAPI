# app/domains/llmtest/service.py
import io
import struct
import zlib
import docx
import olefile
from fastapi import UploadFile, HTTPException, status
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.core.config import settings
from app.domains.llmtest.client import ElasticsearchRagClient
# from app.domains.llmtest.client import OllamaLangChainClient
from app.domains.llmtest.schemas import ChatRequest


class LlmTestService:
    def __init__(self):
        # self.llm_client = OllamaLangChainClient()
        self.rag_client = ElasticsearchRagClient()

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=600,  # bge-m3 모델의 임베딩 해상도를 고려한 청크 크기
            chunk_overlap=60,  # 청크 간 문맥 단절을 막기 위한 중첩 구간
            length_function=len,
            separators=["\n\n", "\n", " ", ""]  # 페이지/문단/문장 단위 순차 안전 분할
        )


    async def get_llm_answer(self, request_data: ChatRequest) -> dict:
        """LangChain 인프라 레이어를 호출하여 결과 스펙을 바인딩합니다."""

        # 1. 비동기 체인 구동
        refined_reply = await self.llm_client.chat_async(request_data.prompt)

        # 2. 도메인 표준 아웃풋 스펙 멀티플렉싱
        return {
            # "model": "exaone3.5",
            "model": settings.OLLAMA_MODEL_NAME,
            "reply": refined_reply
        }
    async def ingest_knowledge_base(self, texts: list[str]) -> dict:
        """사내 매뉴얼이나 참조 문서를 엘라스틱서치 벡터에 적재하는 비즈니스 로직"""
        await self.rag_client.add_documents_async(texts=texts)
        return {"status": "success", "inserted_count": len(texts)}

    async def get_rag_answer(self, request_data: ChatRequest) -> dict:
        """컨텍스트 기반 검색 증강 생성(RAG) 실행"""
        rag_reply = await self.rag_client.ask_rag_async(request_data.prompt)
        return {
            "model": "exaone3.5 + elasticsearch_rag",
            "reply": rag_reply
        }
    async def process_and_ingest_pdf(self, file: UploadFile) -> dict:
        """
        FastAPI UploadFile 스트림 수신 ➡️ 인메모리 PDF 텍스트 추출 ➡️
        문서 청킹(Split) ➡️ 메타데이터 매핑 ➡️ 벡터 스토어 적재 파이프라인
        """
        # 1. 파일 비동기 스트림 읽기 (메모리 적재)
        file_bytes = await file.read()

        # 2. 인메모리 바이너리 스트림 변환 및 pypdf 파싱 (Windows 디스크 I/O Lock 원천 차단)
        pdf_reader = PdfReader(io.BytesIO(file_bytes))

        raw_documents = []
        for page_idx, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if not page_text or not page_text.strip():
                continue

            # 3. 페이지별 로우 데이터 Document 객체화 및 추적용 메타데이터 주입
            raw_documents.append(Document(
                page_content=page_text,
                metadata={
                    "source": file.filename,
                    "page": page_idx + 1  # 1부터 시작하는 페이지 번호
                }
            ))

        # 4. 텍스트 분할기(Text Splitter) 구동하여 청킹 수행
        # 이 과정을 거쳐야 검색 시 엉뚱한 페이지 전체가 아닌 '정확한 단락'만 스코어링되어 LLM에 인입됩니다.
        split_documents = self.text_splitter.split_documents(raw_documents)

        # 5. 인프라 레이어를 통해 엘라스틱서치에 최종 적재
        await self.rag_client.add_langchain_documents_async(split_documents)

        return {
            "status": "success",
            "filename": file.filename,
            "total_pages": len(pdf_reader.pages),
            "chunked_count": len(split_documents)
        }

    def _extract_text_from_hwp(self, file_bytes: bytes) -> str:
        """
        [Core Parser] 한컴오피스 없이 어디서나 동작하는
        순수 파이썬 로우 레벨 HWP5 텍스트 스크래핑 엔진
        """
        try:
            ole = olefile.OleFileIO(io.BytesIO(file_bytes))
            dirs = ole.listdir()

            # HWP 무결성 검증 (정상 문서 포맷 체크)
            if ["FileHeader"] not in dirs:
                raise ValueError("올바른 HWP(아래아한글) 파일 포맷이 아닙니다.")

            # 1. 파일의 zlib 압축 여부 판단 (36번째 바이트 상태 비트 마스킹)
            header = ole.openstream("FileHeader")
            header_data = header.read()
            is_compressed = (header_data[36] & 1) == 1

            # 2. 실제 본문 텍스트 데이터가 들어있는 섹션(Section) 목록 정렬 추적
            section_names = []
            for d in dirs:
                if d[0] == "BodyText" and d[1].startswith("Section"):
                    section_names.append(int(d[1][len("Section"):]))

            sections = [f"BodyText/Section{x}" for x in sorted(section_names)]

            full_text = ""
            # 3. 각 구역 섹션별 바이너리 언팩킹 시전
            for section in sections:
                stream = ole.openstream(section)
                data = stream.read()

                # 압축 문서인 경우 wbits=-15 (Raw Deflate) 속성으로 복원
                unpacked_data = zlib.decompress(data, -15) if is_compressed else data

                # 4. HWP 바이너리 구조에서 글자(HWPTAG_TXT: 67) 레코드 정보만 편취
                size = len(unpacked_data)
                i = 0
                while i < size:
                    header_bytes = struct.unpack_from("<I", unpacked_data, i)[0]
                    rec_type = header_bytes & 0x3FF  # 하위 10비트: 레코드 태그 ID
                    rec_len = (header_bytes >> 20) & 0xFFF  # 상위 12비트: 레코드 길이

                    if rec_type == 67:  # ✨ 67번이 순수 글자 데이터를 보관하는 내부 스펙 코드
                        text_bytes = unpacked_data[i + 4: i + 4 + rec_len]
                        full_text += text_bytes.decode('utf-16le', errors='ignore') + "\n"

                    i += 4 + rec_len  # 다음 레코드로 오프셋 이동

            return full_text
        except Exception as e:
            raise RuntimeError(f"아래아한글 바이너리 디코딩 에러: {str(e)}")

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """
        [Core Parser] python-docx를 활용하여
        본문 단락(Paragraph)과 표(Table) 데이터를 완벽히 추출하는 인메모리 엔진
        """
        try:
            # 인메모리 바이너리 스트림으로 Word 문서 로드
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []

            # 1. 일반 단락 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # 2. [★실무 필수] 표(Table) 내부 데이터 추출
            # 표 안의 데이터를 빼먹으면 기업 매뉴얼 RAG 정합성이 완전히 깨집니다.
            for table in doc.tables:
                for row in table.rows:
                    # 행 단위로 셀 데이터를 탭(\t) 분할하여 문맥 보존
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)
        except Exception as e:
            raise RuntimeError(f"Word 파일(.docx) 디코딩 에러: {str(e)}")

    # async def process_and_ingest_hwp(self, file: UploadFile) -> dict:
    #     """HWP 지식 스트림 인메모리 수신 ➡️ 파싱 ➡️ 청킹 ➡️ ES 벡터 인덱싱 토탈 파이프라인"""
    #     file_bytes = await file.read()
    #
    #     # 1. 커스텀 OLE 파서 구동하여 로우 텍스트 스크래핑
    #     hwp_text = self._extract_text_from_hwp(file_bytes)
    #
    #     if not hwp_text.strip():
    #         return {"status": "skipped", "message": "추출된 유효 문자열이 존재하지 않습니다."}
    #
    #     # 2. 랑체인 표준 스펙 Document 객체 어댑팅
    #     raw_document = Document(
    #         page_content=hwp_text,
    #         metadata={
    #             "source": file.filename,
    #             "page": "Document"  # HWP는 페이지가 유동적이므로 단일 문서 스코프로 래핑
    #         }
    #     )
    #
    #     # 3. 분할기로 600자씩 조각화 수행
    #     split_documents = self.text_splitter.split_documents([raw_document])
    #
    #     # 4. 인프라 클라이언트 레이어를 통해 Elasticsearch 벡터 적재 실행
    #     await self.rag_client.add_langchain_documents_async(split_documents)
    #
    #     return {
    #         "status": "success",
    #         "filename": file.filename,
    #         "chunked_count": len(split_documents)
    #     }

    async def ingest_document(self, file: UploadFile) -> dict:
        """
        [Unified Entrance] 단일 진입점 API.
        확장자를 판별하여 적절한 서브 시스템 파서로 동적 라우팅을 수행합니다.
        """
        filename = file.filename
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        file_bytes = await file.read()

        raw_text = ""
        metadata_page = "Document"

        # 파일 확장자 기반 팩토리 디스패처 구동
        if ext == "pdf":
            # PDF는 페이지별 루프가 필요하므로 내부 파싱 처리를 이관하거나 직접 수행
            from pypdf import PdfReader
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            # 가독성을 위해 간결화된 인메모리 PDF 루프 조합
            raw_text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

        elif ext == "docx":
            raw_text = self._extract_text_from_docx(file_bytes)

        elif ext == "hwp":
            raw_text = self._extract_text_from_hwp(file_bytes)

        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"지원하지 않는 파일 포맷(.{ext})입니다. PDF, DOCX, HWP 파일만 통합 수용합니다."
            )

        if not raw_text.strip():
            return {"status": "skipped", "filename": filename, "reason": "추출된 유효 텍스트가 비어있습니다."}

        # 2. 랑체인 표준 스펙 도큐먼트 래핑 (다형성 확보)
        from langchain_core.documents import Document
        raw_document = Document(
            page_content=raw_text,
            metadata={"source": filename, "type": ext}
        )

        # 3. 통합 분할기(Text Splitter) 가동
        split_documents = self.text_splitter.split_documents([raw_document])

        # 4. 통합 벡터 인덱싱
        await self.rag_client.add_langchain_documents_async(split_documents)

        return {
            "status": "success",
            "filename": filename,
            "format": ext,
            "chunked_count": len(split_documents)
        }